# -*- coding: utf-8 -*-
"""
Created on Thu Aug  1 10:38:55 2019

@author: GuoOu
"""

import docx

"""
filename = 'Decoupling Layer.docx'
doc = docx.Document(filename)
print(len(doc.paragraphs))
for paragraph in doc.paragraphs:
    print(paragraph.text)
#    for run in paragraph.runs:
#        print(run.text)
"""

#create a document
doc = docx.Document()
doc.add_heading('Header 0', 0)
par_one = doc.add_paragraph('')
run = par_one.add_run('Hello world ')
run = par_one.add_run('Another Hello world')
run.underline = True
run.italic = True
run.add_break(docx.text.run.WD_BREAK.PAGE)
par = doc.add_paragraph('')
run = par_one.add_run('Hello again')
doc.save('helloworld.docx')
